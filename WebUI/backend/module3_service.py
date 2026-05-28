"""WebUI Module-3 Service — GraphRAG 任务调度与查询业务逻辑。

本模块从 main.py 中拆离，承担 Module-3 的全部业务逻辑：
  · 路径常量与运行时状态检查
  · 文件名分配（archive_N.txt）
  · GraphRAG 更新/索引任务队列与 worker 线程
  · GraphRAG 查询（全局 / 本地 / 漂移 / 基础）—— 异步化（W-3）
  · 图谱数据读取（parquet → Cytoscape JSON）

调用方（main.py / 路由层）只需 import 本模块函数，无需了解内部实现。
"""
from __future__ import annotations

import asyncio
import io
import json
import os
import queue
import re
import shutil
import subprocess
import tempfile
import threading
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from fastapi import HTTPException

# ---------------------------------------------------------------------------
# 路径常量
# ---------------------------------------------------------------------------

WORKSPACE_ROOT: Path = Path(__file__).resolve().parents[2]
MODULE3_ROOT: Path = WORKSPACE_ROOT / "Module-3"
INPUT_DIR: Path = MODULE3_ROOT / "input"
INPUT_SOURCES_DIR: Path = MODULE3_ROOT / "input_sources"
SETTINGS_FILE: Path = MODULE3_ROOT / "settings.yaml"
MODULE3_CONFIG_FILE: Path = MODULE3_ROOT / "config.yaml"
ENV_FILE: Path = MODULE3_ROOT / ".env"
ROOT_ENV_FILE: Path = WORKSPACE_ROOT / ".env"
OUTPUT_DIR: Path = MODULE3_ROOT / "output"
UPDATE_OUTPUT_DIR: Path = MODULE3_ROOT / "update_output"
os.environ.setdefault("TRAFFIC_GRAPHRAG_REPORT_DIR", str((WORKSPACE_ROOT / "logs" / "graphrag").resolve()))
_RUNTIME_LOG_DIR_FROM_ENV = str(os.getenv("TRAFFIC_RUNTIME_LOG_DIR") or "").strip()
RUNTIME_LOG_DIR: Path = (
    Path(_RUNTIME_LOG_DIR_FROM_ENV).resolve()
    if _RUNTIME_LOG_DIR_FROM_ENV
    else (WORKSPACE_ROOT / "WebUI" / "backend" / "runtime_logs")
)

SUPPORTED_UPLOAD_SUFFIXES = {".txt", ".doc", ".docx", ".pdf"}

WORKFLOW_STAGE_ORDER = [
    "load_input_documents",
    "create_base_text_units",
    "create_final_documents",
    "extract_graph",
    "finalize_graph",
    "extract_covariates",
    "create_communities",
    "create_final_text_units",
    "create_community_reports",
    "generate_text_embeddings",
]

WORKFLOW_STAGE_LABELS = {
    "load_input_documents": "加载输入文档",
    "create_base_text_units": "生成基础文本单元",
    "create_final_documents": "生成最终文档集",
    "extract_graph": "抽取图谱结构",
    "finalize_graph": "图谱整理收敛",
    "extract_covariates": "抽取协变量",
    "create_communities": "构建社区",
    "create_final_text_units": "生成最终文本单元",
    "create_community_reports": "生成社区报告",
    "generate_text_embeddings": "生成文本向量",
}

START_WORKFLOW_PATTERN = re.compile(r"Starting workflow:\s*([A-Za-z0-9_]+)")
COMPLETE_WORKFLOW_PATTERN = re.compile(r"Workflow complete:\s*([A-Za-z0-9_]+)")
QUERY_CITATION_TAG_PATTERN = re.compile(r"\s*\[(?:Data|LLM)\s*:[^\]]*\]\s*", re.IGNORECASE)

QUERY_RESPONSE_STYLE_HINT = (
    "Use Simplified Chinese only. "
    "Do not output any citation or provenance tags such as [Data: ...] or [LLM: verify]."
)


def _is_placeholder_api_key(value: str | None) -> bool:
    text = str(value or "").strip()
    if not text:
        return True
    upper = text.upper()
    if "YOUR_API_KEY" in upper:
        return True
    return text in {"<YOUR_API_KEY>", "YOUR_API_KEY", "<API_KEY>", "API_KEY"}


def _read_env_file_value(path: Path, name: str) -> str | None:
    if not path.exists():
        return None
    try:
        for raw_line in path.read_text(encoding="utf-8").splitlines():
            line = raw_line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, raw_val = line.split("=", 1)
            if key.strip() != name:
                continue
            parsed = raw_val.strip().strip('"').strip("'")
            if parsed and not _is_placeholder_api_key(parsed):
                return parsed
    except Exception:
        return None
    return None


def _read_simple_yaml_scalar_map(path: Path) -> dict[str, str]:
    pairs: dict[str, str] = {}
    try:
        for raw_line in path.read_text(encoding="utf-8").splitlines():
            line = raw_line.strip()
            if not line or line.startswith("#"):
                continue
            if not line.startswith(tuple("ABCDEFGHIJKLMNOPQRSTUVWXYZ_")):
                continue
            if ":" not in line:
                continue
            key, raw_val = line.split(":", 1)
            key = key.strip()
            if not key:
                continue

            value = raw_val.strip()
            if value and not (value.startswith('"') or value.startswith("'")):
                idx = value.find(" #")
                if idx >= 0:
                    value = value[:idx].rstrip()
                if value.startswith("#"):
                    value = ""

            parsed = value.strip().strip('"').strip("'")
            if parsed:
                pairs[key] = parsed
    except Exception:
        return {}
    return pairs


def _read_module3_config_value(name: str) -> str | None:
    if not MODULE3_CONFIG_FILE.exists():
        return None
    pairs = _read_simple_yaml_scalar_map(MODULE3_CONFIG_FILE)
    val = pairs.get(name)
    if val and str(val).strip():
        return str(val).strip()
    return None


def _parse_bool_text(raw: str | None, default: bool) -> bool:
    if raw is None:
        return default
    text = str(raw).strip().lower()
    if text in {"1", "true", "yes", "on"}:
        return True
    if text in {"0", "false", "no", "off"}:
        return False
    return default


def _module3_enable_thinking() -> bool:
    env_raw = os.getenv("MODULE3_ENABLE_THINKING")
    if env_raw and str(env_raw).strip():
        return _parse_bool_text(env_raw, False)

    cfg_raw = _read_module3_config_value("MODULE3_ENABLE_THINKING")
    if cfg_raw is not None:
        return _parse_bool_text(cfg_raw, False)

    return False


def _bootstrap_module3_api_env() -> None:
    # 优先使用已注入的有效环境变量；如未注入，则回退项目根 .env 的 API_KEY。
    base_key = None
    for env_name in ("GRAPHRAG_API_KEY", "DASHSCOPE_API_KEY", "BAILIAN_API_KEY", "API_KEY"):
        value = str(os.getenv(env_name) or "").strip()
        if value and not _is_placeholder_api_key(value):
            base_key = value
            break

    if base_key is None:
        base_key = _read_env_file_value(ROOT_ENV_FILE, "API_KEY")

    if base_key:
        os.environ.setdefault("API_KEY", base_key)
        os.environ.setdefault("GRAPHRAG_API_KEY", base_key)
        os.environ.setdefault("DASHSCOPE_API_KEY", base_key)
        os.environ.setdefault("BAILIAN_API_KEY", base_key)

    thinking_enabled = _module3_enable_thinking()
    os.environ["MODULE3_ENABLE_THINKING"] = "true" if thinking_enabled else "false"


_bootstrap_module3_api_env()

# ---------------------------------------------------------------------------
# 任务队列全局状态
# ---------------------------------------------------------------------------

_update_queue: queue.Queue[str] = queue.Queue()
_tasks: dict[str, dict[str, Any]] = {}
_tasks_lock = threading.Lock()
_name_lock = threading.Lock()
_worker_started = False
_sequence = 0

# ---------------------------------------------------------------------------
# 工具函数
# ---------------------------------------------------------------------------

def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _next_sequence() -> int:
    global _sequence
    _sequence += 1
    return _sequence


def _tail_text(text: str, lines: int = 20) -> str:
    pieces = [l for l in text.splitlines() if l.strip()]
    return "\n".join(pieces[-lines:])


def _effective_query_response_type(response_type: str) -> str:
    base = str(response_type or "").strip() or "Multiple Paragraphs"
    hint_lower = QUERY_RESPONSE_STYLE_HINT.lower()
    if hint_lower in base.lower():
        return base
    return f"{base}. {QUERY_RESPONSE_STYLE_HINT}"


def _sanitize_query_answer(answer: str) -> str:
    text = str(answer or "")
    text = QUERY_CITATION_TAG_PATTERN.sub(" ", text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()


def _decode_console_bytes(raw: bytes) -> str:
    if not raw:
        return ""
    for enc in ("utf-8-sig", "utf-8", "gb18030", "cp936"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _sanitize_graphrag_text(text: str) -> str:
    # 某些文档提取链路会产生 U+FFFD（�），该字符在 Windows GBK 输出链路下可能触发编码异常。
    cleaned = str(text or "").replace("\ufffd", " ")
    cleaned = re.sub(r"[ \t]+\n", "\n", cleaned)
    return cleaned.strip()


def _rel(path: Path) -> str:
    return str(path.relative_to(WORKSPACE_ROOT)).replace("\\", "/")


def _new_workflow_progress() -> dict[str, Any]:
    stages = [
        {
            "key": key,
            "label": WORKFLOW_STAGE_LABELS.get(key, key),
            "status": "pending",
            "started_at": None,
            "finished_at": None,
        }
        for key in WORKFLOW_STAGE_ORDER
    ]
    return {
        "current_stage": None,
        "current_stage_label": None,
        "completed_count": 0,
        "total_count": len(stages),
        "percentage": 0.0,
        "stages": stages,
        "last_line": None,
        "events": [],
    }


def _refresh_workflow_progress(progress: dict[str, Any]) -> None:
    stages = progress.get("stages") or []
    total = len(stages)
    completed = sum(1 for s in stages if s.get("status") == "completed")
    running = any(s.get("status") == "running" for s in stages)
    pct = (completed / total * 100.0) if total else 0.0
    if running and completed < total and total > 0:
        pct += 50.0 / total
    progress["completed_count"] = completed
    progress["total_count"] = total
    progress["percentage"] = round(min(100.0, pct), 1)


def _update_workflow_progress(task_id: str, mutator: Any) -> None:
    with _tasks_lock:
        task = _tasks.get(task_id)
        if not task:
            return
        progress = task.get("workflow_progress")
        if not isinstance(progress, dict):
            progress = _new_workflow_progress()
            task["workflow_progress"] = progress
        mutator(progress)


def _append_workflow_event(progress: dict[str, Any], event: dict[str, Any]) -> None:
    events = progress.get("events")
    if not isinstance(events, list):
        events = []
        progress["events"] = events
    events.append(event)
    if len(events) > 60:
        del events[:-60]


def _mark_workflow_stage_running(task_id: str, stage_key: str, raw_line: str) -> None:
    if stage_key not in WORKFLOW_STAGE_ORDER:
        return

    def _mutator(progress: dict[str, Any]) -> None:
        now = _now_iso()
        stages = progress.get("stages") or []
        for stage in stages:
            if stage.get("status") == "running" and stage.get("key") != stage_key:
                stage["status"] = "pending"
                stage["started_at"] = None
        for stage in stages:
            if stage.get("key") == stage_key:
                if stage.get("status") != "completed":
                    stage["status"] = "running"
                    stage["started_at"] = stage.get("started_at") or now
                progress["current_stage"] = stage_key
                progress["current_stage_label"] = stage.get("label")
                break
        progress["last_line"] = raw_line
        _append_workflow_event(progress, {"type": "start", "stage": stage_key, "time": now, "line": raw_line})
        _refresh_workflow_progress(progress)

    _update_workflow_progress(task_id, _mutator)


def _mark_workflow_stage_completed(task_id: str, stage_key: str, raw_line: str) -> None:
    if stage_key not in WORKFLOW_STAGE_ORDER:
        return

    def _mutator(progress: dict[str, Any]) -> None:
        now = _now_iso()
        stages = progress.get("stages") or []
        for idx, stage in enumerate(stages):
            if stage.get("key") != stage_key:
                continue
            stage["status"] = "completed"
            stage["started_at"] = stage.get("started_at") or now
            stage["finished_at"] = now
            next_stage = stages[idx + 1] if idx + 1 < len(stages) else None
            if next_stage is None:
                progress["current_stage"] = None
                progress["current_stage_label"] = None
            else:
                progress["current_stage"] = next_stage.get("key")
                progress["current_stage_label"] = next_stage.get("label")
            break
        progress["last_line"] = raw_line
        _append_workflow_event(progress, {"type": "complete", "stage": stage_key, "time": now, "line": raw_line})
        _refresh_workflow_progress(progress)

    _update_workflow_progress(task_id, _mutator)


def _mark_workflow_failed(task_id: str, reason: str) -> None:
    def _mutator(progress: dict[str, Any]) -> None:
        now = _now_iso()
        stages = progress.get("stages") or []
        running_stage = None
        for stage in stages:
            if stage.get("status") == "running":
                running_stage = stage
                break
        if running_stage is not None:
            running_stage["status"] = "failed"
            running_stage["started_at"] = running_stage.get("started_at") or now
            running_stage["finished_at"] = now
            progress["current_stage"] = running_stage.get("key")
            progress["current_stage_label"] = running_stage.get("label")
        else:
            # workflow 尚未开始即失败（例如模型配置预检失败）
            progress["current_stage"] = None
            progress["current_stage_label"] = "运行前检查"
        progress["last_line"] = reason
        _append_workflow_event(progress, {"type": "failed", "time": now, "line": reason})
        _refresh_workflow_progress(progress)

    _update_workflow_progress(task_id, _mutator)


def _mark_workflow_success(task_id: str) -> None:
    def _mutator(progress: dict[str, Any]) -> None:
        now = _now_iso()
        stages = progress.get("stages") or []
        for stage in stages:
            if stage.get("status") in {"pending", "running"}:
                stage["status"] = "completed"
                stage["started_at"] = stage.get("started_at") or now
                stage["finished_at"] = now
        progress["current_stage"] = None
        progress["current_stage_label"] = None
        _refresh_workflow_progress(progress)

    _update_workflow_progress(task_id, _mutator)


def _update_workflow_from_line(task_id: str, line: str) -> None:
    if not line:
        return
    start_match = START_WORKFLOW_PATTERN.search(line)
    if start_match:
        _mark_workflow_stage_running(task_id, start_match.group(1).strip(), line)
        return
    complete_match = COMPLETE_WORKFLOW_PATTERN.search(line)
    if complete_match:
        _mark_workflow_stage_completed(task_id, complete_match.group(1).strip(), line)


# ---------------------------------------------------------------------------
# 路径状态校验
# ---------------------------------------------------------------------------

def required_paths_status() -> dict[str, bool]:
    return {
        "module3_root": MODULE3_ROOT.exists(),
        "input_dir": INPUT_DIR.exists(),
        "input_sources_dir": INPUT_SOURCES_DIR.exists(),
        "settings_yaml": SETTINGS_FILE.exists(),
        "env_file": ENV_FILE.exists(),
    }


def validate_runtime_ready() -> None:
    # input 目录允许自动恢复，避免误删后上传接口直接失败。
    try:
        INPUT_DIR.mkdir(parents=True, exist_ok=True)
        INPUT_SOURCES_DIR.mkdir(parents=True, exist_ok=True)
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail={
                "message": "Module-3 input directories cannot be created.",
                "error": str(exc),
                "input_dir": str(INPUT_DIR),
                "input_sources_dir": str(INPUT_SOURCES_DIR),
            },
        ) from exc

    paths = required_paths_status()
    missing = [k for k, ok in paths.items() if not ok]
    if missing:
        raise HTTPException(
            status_code=500,
            detail={
                "message": "Module-3 runtime paths are missing.",
                "missing": missing,
                "paths": {k: str(v) for k, v in {
                    "module3_root": MODULE3_ROOT,
                    "input_dir": INPUT_DIR,
                    "input_sources_dir": INPUT_SOURCES_DIR,
                    "settings_yaml": SETTINGS_FILE,
                    "env_file": ENV_FILE,
                }.items()},
            },
        )


# ---------------------------------------------------------------------------
# 文件名分配
# ---------------------------------------------------------------------------

def get_name_lock() -> threading.Lock:
    return _name_lock


def normalize_uploaded_filename(filename: str) -> str | None:
    raw = str(filename or "").strip()
    if not raw:
        return None
    # 仅保留最终文件名，拦截路径穿越与目录注入。
    name = Path(raw).name.strip()
    if not name or name in {".", ".."}:
        return None
    if any(token in name for token in ("/", "\\", "\x00")):
        return None
    return name


def sanitize_txt_filename_or_raise(filename: str) -> str:
    safe = normalize_uploaded_filename(filename)
    if not safe:
        raise HTTPException(status_code=400, detail=f"Invalid filename: {filename}")
    if not safe.lower().endswith(".txt"):
        raise HTTPException(status_code=400, detail=f"Only .txt files are allowed: {filename}")
    return safe


def sanitize_supported_document_filename_or_raise(filename: str) -> str:
    safe = normalize_uploaded_filename(filename)
    if not safe:
        raise HTTPException(status_code=400, detail=f"Invalid filename: {filename}")
    suffix = Path(safe).suffix.lower()
    if suffix not in SUPPORTED_UPLOAD_SUFFIXES:
        allowed = ", ".join(sorted(SUPPORTED_UPLOAD_SUFFIXES))
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {filename}. Allowed: {allowed}")
    return safe


def resolve_input_text_path_or_raise(file_name: str) -> Path:
    safe = sanitize_txt_filename_or_raise(file_name)
    input_root = INPUT_DIR.resolve()
    target = (INPUT_DIR / safe).resolve()
    if target.parent != input_root:
        raise HTTPException(status_code=400, detail="Invalid file path.")
    return target


def resolve_graph_input_text_path_from_source_name(file_name: str) -> Path:
    safe = sanitize_supported_document_filename_or_raise(file_name)
    input_root = INPUT_DIR.resolve()
    target = (INPUT_DIR / f"{safe}.txt").resolve()
    if target.parent != input_root:
        raise HTTPException(status_code=400, detail="Invalid file path.")
    return target


def _decode_text_bytes(raw: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "gb18030", "cp936"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _extract_txt_text_from_bytes(data: bytes) -> str:
    return _decode_text_bytes(data)


def _extract_docx_text_from_bytes(data: bytes, file_name: str) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"python-docx is required to parse .docx files ({file_name}): {exc}",
        ) from exc

    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Failed to parse .docx file {file_name}: {exc}") from exc

    parts: list[str] = []
    for p in doc.paragraphs:
        text = str(p.text or "").strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [str(cell.text or "").strip() for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _extract_pdf_text_from_bytes(data: bytes, file_name: str) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"pypdf is required to parse .pdf files ({file_name}): {exc}",
        ) from exc

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Failed to parse .pdf file {file_name}: {exc}") from exc

    parts: list[str] = []
    for page in reader.pages:
        text = str(page.extract_text() or "").strip()
        if text:
            parts.append(text)
    return "\n\n".join(parts).strip()


def _extract_doc_text_with_command(data: bytes, tool: str) -> str:
    tool_path = shutil.which(tool)
    if not tool_path:
        return ""
    temp_path: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(data)
            temp_path = Path(tmp.name)
        result = subprocess.run([tool_path, str(temp_path)], capture_output=True, text=False, check=False)
        if result.returncode != 0:
            return ""
        return _decode_console_bytes(result.stdout or b"").strip()
    except Exception:
        return ""
    finally:
        if temp_path is not None and temp_path.exists():
            try:
                temp_path.unlink()
            except Exception:
                pass


def _extract_doc_text_from_bytes(data: bytes, file_name: str) -> str:
    # 优先尝试 textract（若环境具备），其次尝试 antiword/catdoc 命令行工具。
    try:
        import textract  # type: ignore
    except Exception:
        textract = None  # type: ignore

    if textract is not None:
        temp_path: Path | None = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
                tmp.write(data)
                temp_path = Path(tmp.name)
            raw = textract.process(str(temp_path))
            decoded = _decode_text_bytes(bytes(raw)).strip()
            if decoded:
                return decoded
        except Exception:
            pass
        finally:
            if temp_path is not None and temp_path.exists():
                try:
                    temp_path.unlink()
                except Exception:
                    pass

    for tool in ("antiword", "catdoc"):
        text = _extract_doc_text_with_command(data, tool)
        if text:
            return text

    raise HTTPException(
        status_code=400,
        detail=(
            f"Unable to extract text from .doc file: {file_name}. "
            "Please convert it to .docx/.pdf, or install antiword/catdoc (or textract runtime support)."
        ),
    )


def extract_text_from_uploaded_document(file_name: str, data: bytes) -> str:
    suffix = Path(file_name).suffix.lower()
    if suffix == ".txt":
        text = _extract_txt_text_from_bytes(data)
    elif suffix == ".docx":
        text = _extract_docx_text_from_bytes(data, file_name)
    elif suffix == ".pdf":
        text = _extract_pdf_text_from_bytes(data, file_name)
    elif suffix == ".doc":
        text = _extract_doc_text_from_bytes(data, file_name)
    else:
        allowed = ", ".join(sorted(SUPPORTED_UPLOAD_SUFFIXES))
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_name}. Allowed: {allowed}")

    normalized = str(text or "").replace("\r", "").strip()
    if not normalized:
        raise HTTPException(status_code=400, detail=f"No extractable text content found in file: {file_name}")
    return normalized


def save_uploaded_document_for_graphrag(file_name: str, data: bytes) -> dict[str, Any]:
    safe_name = sanitize_supported_document_filename_or_raise(file_name)
    text_path = resolve_graph_input_text_path_from_source_name(safe_name)

    extracted_text = extract_text_from_uploaded_document(safe_name, data)
    extracted_text = _sanitize_graphrag_text(extracted_text)
    INPUT_DIR.mkdir(parents=True, exist_ok=True)

    text_path.write_text(extracted_text, encoding="utf-8")

    text_stat = text_path.stat()
    return {
        "original_name": safe_name,
        "stored_name": text_path.name,
        "stored_path": _rel(text_path),
        "size_bytes": int(text_stat.st_size),
        "graphrag_text_name": text_path.name,
        "graphrag_text_path": _rel(text_path),
        "extracted_chars": len(extracted_text),
    }


def list_uploaded_documents() -> list[dict[str, Any]]:
    docs: list[dict[str, Any]] = []
    INPUT_DIR.mkdir(parents=True, exist_ok=True)
    for entry in INPUT_DIR.iterdir():
        if not entry.is_file() or entry.suffix.lower() != ".txt":
            continue
        stat = entry.stat()
        docs.append({
            "file_name": entry.name,
            "size_bytes": int(stat.st_size),
            "modified_at": datetime.fromtimestamp(stat.st_mtime, timezone.utc).isoformat(),
            "stored_path": _rel(entry),
        })
    docs.sort(key=lambda x: str(x.get("modified_at") or ""), reverse=True)
    return docs


def read_uploaded_document(file_name: str) -> dict[str, Any]:
    # 统一从 input/ 目录读取 .txt 文件
    path = resolve_input_text_path_or_raise(file_name)
    if not path.exists() or not path.is_file():
        raise HTTPException(status_code=404, detail="Document not found.")
    content = path.read_text(encoding="utf-8", errors="replace")

    stat = path.stat()
    return {
        "file_name": path.name,
        "size_bytes": int(stat.st_size),
        "modified_at": datetime.fromtimestamp(stat.st_mtime, timezone.utc).isoformat(),
        "stored_path": _rel(path),
        "content": content,
    }


# ---------------------------------------------------------------------------
# GraphRAG 输出读取（stats）
# ---------------------------------------------------------------------------

def _latest_update_stats() -> dict[str, Any] | None:
    if not UPDATE_OUTPUT_DIR.exists():
        return None
    candidates: list[Path] = []
    for run_dir in UPDATE_OUTPUT_DIR.iterdir():
        if not run_dir.is_dir():
            continue
        sp = run_dir / "delta" / "stats.json"
        if sp.exists():
            candidates.append(sp)
    if not candidates:
        return None
    latest = max(candidates, key=lambda p: p.stat().st_mtime)
    try:
        payload = json.loads(latest.read_text(encoding="utf-8"))
    except Exception as exc:
        return {"stats_path": str(latest), "parse_error": str(exc), "payload": None}
    return {"stats_path": str(latest), "parse_error": None, "payload": payload}


def _latest_output_stats() -> dict[str, Any] | None:
    sp = OUTPUT_DIR / "stats.json"
    if not sp.exists():
        return None
    try:
        payload = json.loads(sp.read_text(encoding="utf-8"))
    except Exception as exc:
        return {"stats_path": str(sp), "parse_error": str(exc), "payload": None}
    return {"stats_path": str(sp), "parse_error": None, "payload": payload}


def reset_knowledge_base() -> dict[str, Any]:
    """物理清空 Module-3 下 cache/input/input_sources/output/update_output。"""
    CACHE_DIR = MODULE3_ROOT / "cache"
    dirs_to_clean = [CACHE_DIR, INPUT_DIR, INPUT_SOURCES_DIR, OUTPUT_DIR, UPDATE_OUTPUT_DIR]
    removed = []
    for d in dirs_to_clean:
        if d.exists():
            shutil.rmtree(d, ignore_errors=True)
            removed.append(str(d.relative_to(WORKSPACE_ROOT)))
    return {"removed": removed}


def _has_output_baseline() -> bool:
    if not OUTPUT_DIR.exists():
        return False
    return any(OUTPUT_DIR.iterdir())


def _assess_update_result(uploaded_count: int, stats_result: dict[str, Any] | None) -> tuple[bool, str, dict[str, Any]]:
    signal: dict[str, Any] = {
        "uploaded_count": uploaded_count, "stats_path": None,
        "num_documents": None, "update_documents": None, "workflow_count": None,
    }
    if stats_result is None:
        if uploaded_count > 0:
            return False, "Update finished but no update_output stats.json was found.", signal
        return True, "No uploaded files and no stats generated.", signal
    signal["stats_path"] = stats_result.get("stats_path")
    if stats_result.get("parse_error"):
        return False, f"Failed to parse GraphRAG stats.json: {stats_result['parse_error']}", signal
    payload = stats_result.get("payload") or {}
    workflows = payload.get("workflows") or {}
    num_documents = int(payload.get("num_documents") or 0)
    update_documents = int(payload.get("update_documents") or 0)
    workflow_count = len(workflows)
    signal.update({"num_documents": num_documents, "update_documents": update_documents, "workflow_count": workflow_count})
    if uploaded_count > 0 and update_documents == 0:
        return False, "GraphRAG returned 0 but no documents were incrementally updated.", signal
    if uploaded_count > 0 and workflow_count == 0:
        return False, "GraphRAG returned 0 but workflow list is empty.", signal
    return True, "Incremental update completed with effective graph signals.", signal


def _assess_index_result(uploaded_count: int, stats_result: dict[str, Any] | None) -> tuple[bool, str, dict[str, Any]]:
    signal: dict[str, Any] = {
        "uploaded_count": uploaded_count, "stats_path": None,
        "num_documents": None, "workflow_count": None,
    }
    if stats_result is None:
        return False, "Index finished but output/stats.json was not found.", signal
    signal["stats_path"] = stats_result.get("stats_path")
    if stats_result.get("parse_error"):
        return False, f"Failed to parse output stats.json: {stats_result['parse_error']}", signal
    payload = stats_result.get("payload") or {}
    workflows = payload.get("workflows") or {}
    num_documents = int(payload.get("num_documents") or 0)
    workflow_count = len(workflows)
    signal.update({"num_documents": num_documents, "workflow_count": workflow_count})
    if uploaded_count > 0 and num_documents == 0:
        return False, "Index completed but no documents were ingested.", signal
    if workflow_count == 0:
        return False, "Index completed but workflows are empty.", signal
    return True, "Bootstrap index completed with effective graph signals.", signal


# ---------------------------------------------------------------------------
# 任务状态管理
# ---------------------------------------------------------------------------

def set_task(task_id: str, patch: dict[str, Any]) -> None:
    with _tasks_lock:
        _tasks[task_id].update(patch)


def enqueue_update_task(trigger: str, uploaded_files: list[str]) -> dict[str, Any]:
    task_id = str(uuid.uuid4())
    task = {
        "task_id": task_id,
        "status": "queued",
        "trigger": trigger,
        "uploaded_files": uploaded_files,
        "created_at": _now_iso(),
        "started_at": None,
        "finished_at": None,
        "return_code": None,
        "error_summary": None,
        "log_file": str((RUNTIME_LOG_DIR / f"update_{task_id}.log").relative_to(WORKSPACE_ROOT)).replace("\\", "/"),
        "sequence": _next_sequence(),
        "workflow_progress": _new_workflow_progress(),
    }
    with _tasks_lock:
        _tasks[task_id] = task
    _update_queue.put(task_id)
    return task


def get_task(task_id: str) -> dict[str, Any] | None:
    with _tasks_lock:
        return _tasks.get(task_id)


def get_recent_tasks(limit: int) -> list[dict[str, Any]]:
    with _tasks_lock:
        return sorted(_tasks.values(), key=lambda x: x["sequence"], reverse=True)[:limit]


def get_queue_position(task_id: str) -> int | None:
    with _tasks_lock:
        task = _tasks.get(task_id)
        if task is None or task["status"] != "queued":
            return None
        queued = sorted((t for t in _tasks.values() if t["status"] == "queued"), key=lambda x: x["sequence"])
    for i, t in enumerate(queued, 1):
        if t["task_id"] == task_id:
            return i
    return None


# ---------------------------------------------------------------------------
# GraphRAG 子进程执行
# ---------------------------------------------------------------------------

def _run_graphrag_update(task_id: str) -> None:
    log_path = RUNTIME_LOG_DIR / f"update_{task_id}.log"
    set_task(task_id, {"status": "running", "started_at": _now_iso()})

    with _tasks_lock:
        current_task = _tasks.get(task_id, {})
        uploaded_count = len(current_task.get("uploaded_files") or [])

    has_baseline = _has_output_baseline()
    if not has_baseline:
        run_mode = "bootstrap_index"
        cmd = ["graphrag", "index", "--root", str(MODULE3_ROOT), "--cache", "--method", "standard", "--verbose"]
    else:
        run_mode = "incremental_update"
        cmd = ["graphrag", "update", "--root", str(MODULE3_ROOT), "--cache", "--method", "standard", "--verbose"]

    set_task(task_id, {"run_mode": run_mode, "command": " ".join(cmd)})

    child_env = os.environ.copy()
    child_env["PYTHONIOENCODING"] = "utf-8"
    child_env["PYTHONUTF8"] = "1"

    try:
        merged_lines: list[str] = []
        process = subprocess.Popen(
            cmd,
            cwd=str(WORKSPACE_ROOT),
            env=child_env,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            encoding="utf-8",
            errors="replace",
            bufsize=1,
        )
        if process.stdout is not None:
            for raw_line in process.stdout:
                line = raw_line.rstrip("\r\n")
                merged_lines.append(line)
                _update_workflow_from_line(task_id, line)
        return_code = int(process.wait())

        merged_output = "\n".join(merged_lines)
        log_path.write_text("# command\n" + " ".join(cmd) + "\n\n# output\n" + merged_output + "\n", encoding="utf-8")

        if return_code == 0:
            _mark_workflow_success(task_id)
            if run_mode == "bootstrap_index":
                ok, message, signal = _assess_index_result(uploaded_count, _latest_output_stats())
            else:
                ok, message, signal = _assess_update_result(uploaded_count, _latest_update_stats())
            if ok:
                set_task(task_id, {"status": "success", "return_code": 0, "finished_at": _now_iso(), "error_summary": None, "knowledge_graph_signal": signal, "summary": message})
            else:
                set_task(task_id, {"status": "failed", "return_code": 0, "finished_at": _now_iso(), "error_summary": message, "knowledge_graph_signal": signal})
        else:
            summary = _tail_text(merged_output) or "GraphRAG update failed."
            _mark_workflow_failed(task_id, summary)
            set_task(task_id, {"status": "failed", "return_code": return_code, "finished_at": _now_iso(), "error_summary": summary})
    except Exception as exc:
        _mark_workflow_failed(task_id, str(exc))
        log_path.write_text(f"Update failed before completion: {exc}\n", encoding="utf-8")
        set_task(task_id, {"status": "failed", "return_code": -1, "finished_at": _now_iso(), "error_summary": str(exc)})


def _worker_loop() -> None:
    while True:
        task_id = _update_queue.get()
        _run_graphrag_update(task_id)
        _update_queue.task_done()


def ensure_worker_started() -> None:
    global _worker_started
    if _worker_started:
        return
    thread = threading.Thread(target=_worker_loop, daemon=True, name="graphrag-update-worker")
    thread.start()
    _worker_started = True


def get_queue_size() -> int:
    return _update_queue.qsize()


# ---------------------------------------------------------------------------
# GraphRAG 查询 —— 异步（W-3）
# ---------------------------------------------------------------------------

def _run_graphrag_query_sync(
    question: str,
    method: str,
    response_type: str,
    community_level: int | None,
    verbose: bool,
) -> dict[str, Any]:
    """同步执行 graphrag query 子进程。由 asyncio.to_thread 在线程池中调用。"""
    query_id = str(uuid.uuid4())
    log_path = RUNTIME_LOG_DIR / f"query_{query_id}.log"
    effective_response_type = _effective_query_response_type(response_type)

    cmd = [
        "graphrag", "query",
        "--root", str(MODULE3_ROOT),
        "--method", method,
        "--response-type", effective_response_type,
    ]
    if community_level is not None:
        cmd.extend(["--community-level", str(community_level)])
    if verbose:
        cmd.append("--verbose")
    cmd.append(question)

    child_env = os.environ.copy()
    child_env["PYTHONIOENCODING"] = "utf-8"
    child_env["PYTHONUTF8"] = "1"

    result = subprocess.run(cmd, cwd=str(WORKSPACE_ROOT), env=child_env, capture_output=True, text=False, check=False)
    stdout = _decode_console_bytes(result.stdout or b"")
    stderr = _decode_console_bytes(result.stderr or b"")
    log_path.write_text("# command\n" + " ".join(cmd) + "\n\n# stdout\n" + stdout + "\n\n# stderr\n" + stderr, encoding="utf-8")

    if result.returncode != 0:
        raise HTTPException(
            status_code=500,
            detail={
                "message": "GraphRAG query failed.",
                "return_code": result.returncode,
                "error_summary": _tail_text(stderr) or _tail_text(stdout),
                "log_file": _rel(log_path),
            },
        )

    answer = _sanitize_query_answer(stdout)
    if not answer:
        answer = stdout.strip()
    if not answer:
        raise HTTPException(
            status_code=500,
            detail={"message": "GraphRAG query returned empty output.", "log_file": _rel(log_path)},
        )

    return {
        "query_id": query_id,
        "method": method,
        "response_type": effective_response_type,
        "question": question,
        "answer": answer,
        "log_file": _rel(log_path),
    }


async def run_graphrag_query_async(
    question: str,
    method: str,
    response_type: str,
    community_level: int | None = None,
    verbose: bool = False,
) -> dict[str, Any]:
    """异步包装：在 asyncio 线程池中运行 GraphRAG 子进程，不阻塞 event loop。"""
    return await asyncio.to_thread(
        _run_graphrag_query_sync,
        question, method, response_type, community_level, verbose,
    )


# ---------------------------------------------------------------------------
# 图谱数据（parquet → Cytoscape JSON）
# ---------------------------------------------------------------------------

def build_graph_data(max_nodes: int, max_edges: int) -> dict[str, Any]:
    entities_path = OUTPUT_DIR / "entities.parquet"
    relationships_path = OUTPUT_DIR / "relationships.parquet"

    if not entities_path.exists() or not relationships_path.exists():
        raise HTTPException(
            status_code=404,
            detail={
                "message": "Graph outputs not found. Please run GraphRAG index/update first.",
                "entities_path": str(entities_path),
                "relationships_path": str(relationships_path),
            },
        )

    try:
        import pyarrow.parquet as pq
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail={"message": "pyarrow is required to read GraphRAG parquet outputs.", "error": str(exc)},
        ) from exc

    ent = pq.read_table(entities_path).to_pydict()
    titles = ent.get("title", [])
    ids = ent.get("id", [])
    types = ent.get("type", [])
    descriptions = ent.get("description", [])
    total_nodes = len(titles)
    take_nodes = min(total_nodes, max_nodes)

    nodes: list[dict[str, Any]] = []
    title_to_id: dict[str, str] = {}
    for i in range(take_nodes):
        title = str(titles[i]) if i < len(titles) and titles[i] is not None else f"node-{i}"
        node_id = str(ids[i]) if i < len(ids) and ids[i] is not None else f"title::{title}"
        node_type = str(types[i]) if i < len(types) and types[i] is not None else "UNKNOWN"
        desc = str(descriptions[i]) if i < len(descriptions) and descriptions[i] is not None else ""
        title_to_id[title] = node_id
        nodes.append({"data": {"id": node_id, "label": title, "type": node_type, "description": desc[:300]}})

    rel = pq.read_table(relationships_path).to_pydict()
    rel_ids = rel.get("id", [])
    sources = rel.get("source", [])
    targets = rel.get("target", [])
    weights = rel.get("weight", [])
    rel_descs = rel.get("description", [])
    total_edges = len(sources)

    edges: list[dict[str, Any]] = []
    for i in range(total_edges):
        if len(edges) >= max_edges:
            break
        src = str(sources[i]) if i < len(sources) and sources[i] is not None else ""
        tgt = str(targets[i]) if i < len(targets) and targets[i] is not None else ""
        if not src or not tgt:
            continue
        src_id = title_to_id.get(src)
        tgt_id = title_to_id.get(tgt)
        if src_id is None or tgt_id is None:
            continue
        edge_id = str(rel_ids[i]) if i < len(rel_ids) and rel_ids[i] is not None else f"edge-{i}"
        weight = float(weights[i]) if i < len(weights) and weights[i] is not None else 1.0
        desc = str(rel_descs[i]) if i < len(rel_descs) and rel_descs[i] is not None else ""
        edges.append({"data": {"id": edge_id, "source": src_id, "target": tgt_id, "weight": weight, "description": desc[:240]}})

    return {
        "nodes": nodes, "edges": edges,
        "meta": {
            "entities_path": _rel(entities_path), "relationships_path": _rel(relationships_path),
            "total_nodes": total_nodes, "total_edges": total_edges,
            "returned_nodes": len(nodes), "returned_edges": len(edges),
            "max_nodes": max_nodes, "max_edges": max_edges,
        },
    }
